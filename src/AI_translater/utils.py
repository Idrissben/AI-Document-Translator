import docx
import tiktoken


# Functions to read and write .docx files
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip() != "":
            full_text.append(para.text)
    return "\n".join(full_text)


def write_docx(text, file_path):
    doc = docx.Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(file_path)


# Token counting function
def count_tokens(text):
    encoding = tiktoken.encoding_for_model("gpt-4.0")
    tokens = encoding.encode(text)
    return len(tokens)
